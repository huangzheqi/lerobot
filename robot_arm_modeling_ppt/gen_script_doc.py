# -*- coding: utf-8 -*-
"""Generate the 3-minute presentation script as .md and .docx (kept in sync)."""
import os, re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))

TITLE = "机械臂抓取任务的动态系统建模——3 分钟汇报稿"
NOTE = ("口播正文约 720 字，正常汇报语速（约 240 字/分钟）3 分钟左右。"
        "【】内为对应 PPT 页面与建议用时，（）内为动作提示，均不必念出。")

SECTIONS = [
    ("【封面 ｜ 约 20 秒】", [
        "各位老师好，我汇报的主题是「机械臂抓取任务的动态系统建模」。"
        "工作分三步：建立动力学方程，变换到状态空间，再求出传递函数，"
        "全程用 MATLAB 实现并验证。（指封面流程图，从左到右带一遍）",
    ]),
    ("【第 1 页 · 动力学建模 ｜ 约 55 秒】", [
        "研究对象是平面二连杆机械臂（指右图）：两个关节由力矩驱动，"
        "末端抓取 0.5 公斤的负载，视为固连在连杆二末端的点质量。",
        "建模用拉格朗日方法：写出包含负载的动能和势能，代入拉格朗日方程，"
        "得到中间高亮的动力学方程（指蓝框）。M(q) 是惯性矩阵，C 是科氏力、"
        "离心力项，G(q) 是重力矩，Fv 是关节摩擦，τ 是控制输入。",
        "关键点是：抓取后 M 和 G 里都多出含负载质量 mp 的项——模型随抓取状态"
        "而改变，这正是要做建模分析的原因。",
    ]),
    ("【第 2 页 · 状态空间变换 ｜ 约 55 秒】", [
        "动力学方程是非线性的，所以第二步在平衡点附近线性化。我们选悬垂抓取位形："
        "两杆竖直向下，此时重力矩为零、平衡输入也为零，分析最干净。",
        "小偏差展开时，速度乘积的 C 项是二阶小量、可以略去，得到这个线性方程"
        "（指公式）；其中 Kg 是重力刚度矩阵，即重力矩对关节角的导数。",
        "再取关节角偏差加角速度组成四维状态向量，输入是力矩、输出是关节角，"
        "就得到标准的状态空间模型：A、B、C、D 都有解析的分块表达式（指底部蓝框），"
        "右侧是代入数值后的矩阵。",
    ]),
    ("【第 3 页 · 传递函数与 MATLAB ｜ 约 55 秒】", [
        "第三步由状态空间求传递函数：G(s) = C(sI−A)⁻¹B + D。左边的 MATLAB 代码"
        "核心只有几行：ss 构建状态空间模型，tf 或 ss2tf 转成传递函数；"
        "系统两输入两输出，得到 2×2 的传递函数矩阵。",
        "右边是运行结果：关节一力矩到关节一角度通道，是一个四阶传递函数（指蓝框）。"
        "四个极点全部位于左半平面，说明抓取平衡位形附近渐近稳定；主导极点阻尼比"
        "约 0.05，呈欠阻尼摆动。直流增益与重力刚度矩阵的逆一致，模型自洽。",
    ]),
    ("【收尾 ｜ 约 15 秒】", [
        "总结：我们打通了「动力学方程—状态空间—传递函数」的完整建模链条，"
        "并经 MATLAB 验证。它是后续 PID、LQR 等控制设计和负载变化分析的基础。"
        "我的汇报到此结束，谢谢大家！",
    ]),
]

QA = [
    ("为什么选悬垂位形作平衡点？",
     "该位形下 G(q0)=0、平衡力矩 τ0=0，结果最简洁；方法本身对任意工作点同样适用，只是 Kg 与 τ0 取值不同。"),
    ("为什么线性化时可以丢掉 C 项？",
     "科氏力/离心力项与关节角速度的乘积（平方）成正比，平衡点处速度为零，故它是偏差量的二阶小量。"),
    ("负载质量变化了怎么办？",
     "M0 与 Kg 随 mp 解析地变化，可代入新参数重新计算模型；也可以据此做鲁棒控制或自适应控制设计。"),
]

# ---------------- markdown ----------------
md = [f"# {TITLE}", "", f"> {NOTE}", ""]
for head, paras in SECTIONS:
    md.append(f"## {head}")
    md.append("")
    for p in paras:
        md.append(p)
        md.append("")
md += ["---", "", "## 附：可能的提问与简答（不计入 3 分钟）", ""]
for q, a in QA:
    md.append(f"- **问：{q}**  \n  答：{a}")
md.append("")
with open(os.path.join(HERE, "汇报稿_3分钟.md"), "w", encoding="utf-8") as f:
    f.write("\n".join(md))
print("saved 汇报稿_3分钟.md")

# ---------------- docx ----------------
NAVY = RGBColor(0x1F, 0x2A, 0x44)
DEEP = RGBColor(0x2E, 0x5E, 0x8C)
GRAY = RGBColor(0x6B, 0x72, 0x80)

doc = Document()
for sec in doc.sections:
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    sec.left_margin = sec.right_margin = Cm(2.4)

def set_ea(run, ea="微软雅黑", latin="Calibri"):
    run.font.name = latin
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), ea)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(TITLE); r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = NAVY; set_ea(r)

p = doc.add_paragraph()
r = p.add_run(NOTE); r.font.size = Pt(9.5); r.font.color.rgb = GRAY; set_ea(r)

for head, paras in SECTIONS:
    p = doc.add_paragraph(); p.space_before = Pt(10)
    r = p.add_run(head); r.font.size = Pt(12.5); r.font.bold = True; r.font.color.rgb = DEEP; set_ea(r)
    for txt in paras:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(21)
        p.paragraph_format.line_spacing = 1.35
        r = p.add_run(txt); r.font.size = Pt(11); r.font.color.rgb = NAVY; set_ea(r)

p = doc.add_paragraph(); p.space_before = Pt(14)
r = p.add_run("附：可能的提问与简答（不计入 3 分钟）")
r.font.size = Pt(12.5); r.font.bold = True; r.font.color.rgb = DEEP; set_ea(r)
for q, a in QA:
    p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.3
    r = p.add_run("问：" + q); r.font.size = Pt(10.5); r.font.bold = True; r.font.color.rgb = NAVY; set_ea(r)
    p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.3
    r = p.add_run("答：" + a); r.font.size = Pt(10.5); r.font.color.rgb = NAVY; set_ea(r)

doc.core_properties.title = TITLE
doc.save(os.path.join(HERE, "汇报稿_3分钟.docx"))
print("saved 汇报稿_3分钟.docx")

total = sum(len(t) for _, ps in SECTIONS for t in ps)
spoken = sum(len(re.sub(r"（[^）]*）", "", t)) for _, ps in SECTIONS for t in ps)
print("总字数:", total, "| 口播字数(去动作提示):", spoken)

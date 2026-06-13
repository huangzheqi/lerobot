# -*- coding: utf-8 -*-
"""生成《基于强化学习的机械臂线性二次型最优控制研究》课程论文 Word 文档。
格式参照《浙江大学学报（工学版）》：中文标题/作者/单位 + 中文摘要、关键词、
中图分类号、文献标志码 + 英文标题/作者/单位 + 英文摘要、关键词。
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CENTER = WD_ALIGN_PARAGRAPH.CENTER
JUSTIFY = WD_ALIGN_PARAGRAPH.JUSTIFY


def set_font(run, cn='宋体', en='Times New Roman', size=10.5, bold=False, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), en)
    rFonts.set(qn('w:hAnsi'), en)
    rFonts.set(qn('w:eastAsia'), cn)


def add_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=None, first_indent=0):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = line
    if first_indent:
        pf.first_line_indent = Pt(first_indent)
    return p


# ---------------- 文本内容 ----------------
TITLE_CN = '基于强化学习的机械臂线性二次型最优控制研究'
AUTHOR_CN = '（请填写姓名）'
AFFIL_CN = '（浙江大学  ____________学院，浙江  杭州  310027）'

ABS_CN = (
    '针对机械臂在模型参数不确定条件下的最优轨迹跟踪控制问题，本文从线性系统理论出发，'
    '研究将线性二次型调节器（LQR）与强化学习相结合的控制方法。首先，对 n 自由度机械臂的'
    '非线性动力学方程在平衡点（或参考轨迹）附近进行线性化，建立其状态空间模型，并对系统的'
    '可控性、可观性与稳定性进行分析；在此基础上，将最优控制问题归结为求解代数 Riccati 方程'
    '（ARE）的线性二次型调节问题。针对实际工程中系统模型参数难以精确获取的情形，引入基于'
    '自适应动态规划（ADP）的策略迭代算法，以在线、数据驱动的方式逼近 Riccati 方程的解，'
    '在无需精确系统模型的条件下学习最优状态反馈增益，并利用 Lyapunov 理论证明闭环系统的'
    '稳定性与迭代算法的收敛性。最后，以二自由度平面机械臂为对象进行数值仿真，将所提强化'
    '学习控制器与模型已知的传统 LQR 控制器进行对比。仿真结果表明，在系统模型未知的条件下，'
    '所提方法能够收敛到接近最优的状态反馈增益，实现稳定、快速的关节轨迹跟踪，从而验证了'
    '线性系统理论与强化学习相结合在机械臂控制中的有效性。'
)
KW_CN = '线性系统理论；线性二次型调节器（LQR）；强化学习；自适应动态规划；机械臂控制；代数 Riccati 方程'
CLC = 'TP242.6'
DOC_CODE = 'A'

TITLE_EN = ('Reinforcement learning based linear quadratic optimal control '
            'of robotic manipulators')
AUTHOR_EN = '(Author Name)'
AFFIL_EN = '(School of ____________, Zhejiang University, Hangzhou 310027, China)'
ABS_EN = (
    'To address the optimal trajectory-tracking control of robotic manipulators under '
    'model uncertainty, this paper, starting from linear system theory, investigates a '
    'reinforcement learning (RL) control method that integrates the linear quadratic '
    'regulator (LQR). First, the nonlinear dynamics of an n-degree-of-freedom manipulator '
    'are linearized about an equilibrium point (or a reference trajectory) to obtain a '
    'state-space model, whose controllability, observability and stability are analyzed. '
    'The optimal control problem is then reformulated as an LQR problem whose solution '
    'requires solving the algebraic Riccati equation (ARE). Considering that exact model '
    'parameters are often unavailable in practice, a policy-iteration algorithm based on '
    'adaptive dynamic programming (ADP) is introduced to approximate the solution of the '
    'ARE in an online, data-driven manner, thereby learning the optimal state-feedback '
    'gain without an exact system model; the closed-loop stability and the convergence of '
    'the iterative algorithm are established by means of Lyapunov theory. Finally, '
    'numerical simulations on a two-degree-of-freedom planar manipulator compare the '
    'proposed RL controller with a conventional model-based LQR controller. The results '
    'show that, without prior knowledge of the model, the proposed method converges to a '
    'near-optimal feedback gain and achieves stable and fast joint trajectory tracking, '
    'verifying the effectiveness of combining linear system theory with reinforcement '
    'learning for manipulator control.'
)
KW_EN = ('linear system theory; linear quadratic regulator (LQR); reinforcement learning; '
         'adaptive dynamic programming; manipulator control; algebraic Riccati equation')

# ---------------- 构建文档 ----------------
doc = Document()

sec = doc.sections[0]
sec.page_height = Cm(29.7)
sec.page_width = Cm(21.0)
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(2.8)
sec.right_margin = Cm(2.8)

normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 中文标题
p = add_para(doc, CENTER, before=6, after=8)
set_font(p.add_run(TITLE_CN), cn='黑体', en='SimHei', size=16, bold=True)

# 中文作者
p = add_para(doc, CENTER, after=2)
set_font(p.add_run(AUTHOR_CN), cn='楷体', size=12)

# 中文单位
p = add_para(doc, CENTER, after=10)
set_font(p.add_run(AFFIL_CN), cn='宋体', size=9)

# 中文摘要
p = add_para(doc, JUSTIFY, after=3, line=1.25)
set_font(p.add_run('摘要：'), cn='黑体', en='SimHei', size=9, bold=True)
set_font(p.add_run(ABS_CN), cn='宋体', size=9)

# 中文关键词
p = add_para(doc, JUSTIFY, after=3, line=1.25)
set_font(p.add_run('关键词：'), cn='黑体', en='SimHei', size=9, bold=True)
set_font(p.add_run(KW_CN), cn='宋体', size=9)

# 中图分类号 + 文献标志码
p = add_para(doc, JUSTIFY, after=12, line=1.25)
set_font(p.add_run('中图分类号：'), cn='黑体', en='SimHei', size=9, bold=True)
set_font(p.add_run(CLC + '　　'), cn='宋体', size=9)
set_font(p.add_run('文献标志码：'), cn='黑体', en='SimHei', size=9, bold=True)
set_font(p.add_run(DOC_CODE), cn='宋体', size=9)

# 英文标题
p = add_para(doc, CENTER, before=4, after=6)
set_font(p.add_run(TITLE_EN), en='Times New Roman', size=12, bold=True)

# 英文作者
p = add_para(doc, CENTER, after=2)
set_font(p.add_run(AUTHOR_EN), en='Times New Roman', size=10.5)

# 英文单位
p = add_para(doc, CENTER, after=10)
set_font(p.add_run(AFFIL_EN), en='Times New Roman', size=9, italic=True)

# 英文摘要
p = add_para(doc, JUSTIFY, after=3, line=1.25)
set_font(p.add_run('Abstract: '), en='Times New Roman', size=9, bold=True)
set_font(p.add_run(ABS_EN), en='Times New Roman', size=9)

# 英文关键词
p = add_para(doc, JUSTIFY, after=3, line=1.25)
set_font(p.add_run('Key words: '), en='Times New Roman', size=9, bold=True)
set_font(p.add_run(KW_EN), en='Times New Roman', size=9)

out = 'course_paper/课程论文_基于强化学习的机械臂线性二次型最优控制研究.docx'
doc.save(out)
print('saved:', out)
